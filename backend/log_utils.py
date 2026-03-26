"""
log_utils.py
────────────
Writes entry records to a Word document.
Each row: Photo | Sr. | Name | Date | Time | Confidence | Status
"""

import os
import json
import base64
import tempfile
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _make_new_doc() -> Document:
    doc = Document()

    # Margins
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin   = Inches(0.7)
    sec.right_margin  = Inches(0.7)

    # Title
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("🛡️  FaceGuard — Entry Log")
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr  = sub.add_run(f"Auto-generated  |  {datetime.now().strftime('%d %B %Y')}")
    sr.font.size  = Pt(9)
    sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # Table
    headers = ["Photo", "Sr.", "Name", "Date", "Time", "Confidence", "Status"]
    table   = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for i, (hdr, cell) in enumerate(zip(headers, hdr_row.cells)):
        cell.text = hdr
        r = cell.paragraphs[0].runs[0]
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "10B981")

    return doc


def log_entry_to_word(
    doc_path     : str,
    name         : str,
    department   : str,
    timestamp    : str,
    confidence   : float = 0.0,
    screenshot_b64: str  = None,
):
    path = Path(doc_path)
    doc  = Document(str(path)) if path.exists() else _make_new_doc()

    table     = doc.tables[-1]
    row_count = len(table.rows)   # header = row 1, so first data row = 2
    row       = table.add_row()

    # ── Col 0: Photo ──────────────────────────────────────────────────────
    cell_photo = row.cells[0]
    cell_photo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if screenshot_b64:
        try:
            raw = base64.b64decode(screenshot_b64.split(",")[-1])
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                tmp.write(raw)
                tmp_path = tmp.name
            run = cell_photo.paragraphs[0].add_run()
            run.add_picture(tmp_path, width=Inches(0.85))
            os.unlink(tmp_path)
        except Exception:
            cell_photo.paragraphs[0].add_run("📷")
    else:
        cell_photo.paragraphs[0].add_run("—")

    # ── Col 1: Sr. ────────────────────────────────────────────────────────
    row.cells[1].text = str(row_count)
    _style(row.cells[1], Pt(9))

    # ── Col 2: Name ───────────────────────────────────────────────────────
    row.cells[2].text = name
    _style(row.cells[2], Pt(9), bold=True, color=RGBColor(0x10, 0xB9, 0x81))

    # ── Col 3: Date ───────────────────────────────────────────────────────
    dt = datetime.strptime(timestamp, "%Y-%m-%d %H:%M:%S")
    row.cells[3].text = dt.strftime("%d %b %Y")
    _style(row.cells[3], Pt(9))

    # ── Col 4: Time ───────────────────────────────────────────────────────
    row.cells[4].text = dt.strftime("%I:%M:%S %p")
    _style(row.cells[4], Pt(9))

    # ── Col 5: Confidence ─────────────────────────────────────────────────
    row.cells[5].text = f"{confidence:.1f}%"
    if confidence >= 80:
        color = RGBColor(0x10, 0xB9, 0x81)
    elif confidence >= 60:
        color = RGBColor(0xF5, 0x9E, 0x0B)
    else:
        color = RGBColor(0xEF, 0x44, 0x44)
    _style(row.cells[5], Pt(9), color=color)

    # ── Col 6: Status ─────────────────────────────────────────────────────
    row.cells[6].text = "✅ Verified"
    _style(row.cells[6], Pt(9))

    # Align all center
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Zebra stripe
    if row_count % 2 == 0:
        for cell in row.cells:
            _set_cell_bg(cell, "F0FDF9")

    doc.save(str(path))

    # Also append to JSON backup
    json_path = path.parent / "entries.json"
    entries   = []
    if json_path.exists():
        try:
            with open(json_path) as f:
                entries = json.load(f)
        except Exception:
            entries = []
    entries.append({
        "name"      : name,
        "department": department,
        "timestamp" : timestamp,
        "confidence": confidence,
    })
    with open(json_path, "w") as f:
        json.dump(entries, f, indent=2)


def _style(cell, size, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = size
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
